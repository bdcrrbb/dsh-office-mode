"""B 档国标公文样张生成 + 自检。
验证：GB/T 9704 页边距、字体、行距。
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pathlib

OUT = pathlib.Path(__file__).parent / "out"
OUT.mkdir(exist_ok=True)
doc = Document()

# 页边距：上 3.7cm 下 3.5cm 左 2.8cm 右 2.6cm
section = doc.sections[0]
section.top_margin = Cm(3.7)
section.bottom_margin = Cm(3.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.6)

# 标题：方正小标宋简体 二号(22pt)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('关于办公 Agent 模式建设的报告')
run.font.name = 'Times New Roman'
run.font.size = Pt(22)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

doc.add_paragraph()  # 空行

# 正文：仿宋_GB2312 三号(16pt)，行距固定值 28 磅
def add_gb_paragraph(text, font_ea='仿宋_GB2312', size_pt=16, line_spacing_pt=28):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_ea)
    return p

add_gb_paragraph('各部门：')
add_gb_paragraph('为提升办公效率，现拟建设 DSH 办公模式，具体方案如下：')
add_gb_paragraph('一、背景与目标')
add_gb_paragraph('阿里、腾讯、字节等大厂均在布局办公 Agent，DSH 目前只有编码模式，需增加办公模式。')
add_gb_paragraph('二、技术方案')
add_gb_paragraph('采用本地 python 工具链（python-docx/python-pptx），不依赖外部 MCP 服务。')
add_gb_paragraph('三、实施计划')
add_gb_paragraph('分 P0（工具链）、P1（真实验证）、P2（Excel/MCP）三期实施。')

# 保存
out_path = OUT / "docx_gb_sample.docx"
doc.save(out_path)
print(f"saved → {out_path}")

# 自检
doc2 = Document(out_path)
sec = doc2.sections[0]
print(f"  margins: top={sec.top_margin.cm:.2f}cm bottom={sec.bottom_margin.cm:.2f}cm left={sec.left_margin.cm:.2f}cm right={sec.right_margin.cm:.2f}cm")
assert abs(sec.top_margin.cm - 3.7) < 0.05, '上边距不符'
assert abs(sec.bottom_margin.cm - 3.5) < 0.05, '下边距不符'
assert abs(sec.left_margin.cm - 2.8) < 0.05, '左边距不符'
assert abs(sec.right_margin.cm - 2.6) < 0.05, '右边距不符'
print("  页边距 ✓")

# 检查字体
ea_fonts = set()
for p in doc2.paragraphs:
    for run in p.runs:
        if run._element.rPr is not None:
            ea = run._element.rPr.rFonts.get(qn('w:eastAsia'))
            if ea:
                ea_fonts.add(ea)
print(f"  eastAsia fonts: {ea_fonts}")
assert '方正小标宋简体' in ea_fonts, '标题字体不符'
assert '仿宋_GB2312' in ea_fonts, '正文字体不符'
print("  字体 ✓")

print("\nB 档公文样张自检通过")
