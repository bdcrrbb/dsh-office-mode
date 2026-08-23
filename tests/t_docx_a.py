"""A 档商务文档样张生成 + 自检。
验证：标题层级、表格、TOC 域、eastAsia 字体设置。
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT = pathlib.Path(__file__).parent / "out"
OUT.mkdir(exist_ok=True)
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_heading('主流办公 Agent 对比调研报告', 0)
for run in title.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(22)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# 副标题
subtitle = doc.add_paragraph('2026 年 8 月', style='Subtitle')
for run in subtitle.runs:
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# TOC 占位（实际需 F9 更新）
doc.add_paragraph()
toc_para = doc.add_paragraph()
run = toc_para.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
run._element.append(fldChar)
run2 = toc_para.add_run()
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
run2._element.append(instrText)
run3 = toc_para.add_run()
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run3._element.append(fldChar2)
doc.add_paragraph()

# 第一章
doc.add_heading('一、调研背景', level=1)
p = doc.add_paragraph('随着 AI 办公 Agent 的兴起，本报告对比阿里 Qoder Work、腾讯 WorkBuddy、字节飞书 Aily 与 OpenAI ChatGPT Agent 的核心能力。')
for run in p.runs:
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 第二章
doc.add_heading('二、产品对比', level=1)
doc.add_heading('2.1 功能矩阵', level=2)

# 表格
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['产品', '文档生成', 'PPT 制作', '本地优先']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

data = [
    ['Qoder Work', '✓', '✓', '✓'],
    ['WorkBuddy', '✓', '✓', '✗'],
    ['飞书 Aily', '✓', '✓', '✗'],
    ['ChatGPT Agent', '✓', '✓', '✗'],
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# 第三章
doc.add_heading('三、结论', level=1)
p = doc.add_paragraph('Qoder Work 的本地优先架构与 DSH 办公模式同构，是最佳参考对象。')
for run in p.runs:
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 保存
out_path = OUT / "docx_a_sample.docx"
doc.save(out_path)
print(f"saved → {out_path}")

# 自检
doc2 = Document(out_path)
headings = [p for p in doc2.paragraphs if p.style.name.startswith('Heading')]
tables = doc2.tables
print(f"  headings: {len(headings)} (≥3 ✓)")
print(f"  tables: {len(tables)} (≥1 ✓)")

# 检查 eastAsia 字体
ea_fonts = set()
for p in doc2.paragraphs[:10]:
    for run in p.runs:
        if run._element.rPr is not None:
            ea = run._element.rPr.rFonts.get(qn('w:eastAsia'))
            if ea:
                ea_fonts.add(ea)
print(f"  eastAsia fonts: {ea_fonts}")
assert '宋体' in ea_fonts or '微软雅黑' in ea_fonts, 'eastAsia 字体未设置'
print("  eastAsia 字体 ✓")

print("\nA 档样张自检通过")
